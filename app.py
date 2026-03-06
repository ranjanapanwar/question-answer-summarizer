"""
MCQ PDF → Answer Sheet Generator

Reads PDFs with questions formatted as:
  Question: N. question text?
  (a) option 1  (b) option 2  (c) option 3  (d) option 4
  Answer: x
  Positive Marks: 1
  Negative Marks: 0

Outputs a clean Q&A document (PDF or DOCX).
"""

import re
import tempfile

import gradio as gr
import pdfplumber
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# PDF text extraction – column-aware
# ---------------------------------------------------------------------------

def extract_text_column_aware(pdf_path: str) -> str:
    """
    Extract text from a multi-column PDF by clustering words on x-axis
    so questions aren't interleaved across columns.
    """
    all_text = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(x_tolerance=3, y_tolerance=3)
            if not words:
                continue

            # Detect column boundaries by clustering word x0 positions
            x_positions = sorted(set(round(w["x0"] / 50) * 50 for w in words))
            # Use page thirds as column boundaries (works well for 2-3 column layouts)
            page_width = page.width
            col_boundaries = [0, page_width / 3, 2 * page_width / 3, page_width]

            def get_column(x0):
                for i in range(len(col_boundaries) - 1):
                    if col_boundaries[i] <= x0 < col_boundaries[i + 1]:
                        return i
                return len(col_boundaries) - 2

            # Group words by (column, row) where row = rounded y-coordinate
            from collections import defaultdict
            rows: dict = defaultdict(list)
            for w in words:
                col = get_column(w["x0"])
                row_y = round(w["top"] / 5) * 5  # 5pt tolerance
                rows[(col, row_y)].append(w)

            # Sort and reconstruct text column by column, top to bottom
            page_lines = []
            for (col, row_y), row_words in sorted(rows.items()):
                line = " ".join(w["text"] for w in sorted(row_words, key=lambda w: w["x0"]))
                page_lines.append(line)

            all_text.append("\n".join(page_lines))

    return "\n".join(all_text)


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------

def parse_qa(text: str) -> list[dict]:
    """
    Parse Q&A pairs from the extracted text.

    Supports formats:
      - 'Question: N. text'  (RIMC / standard test format)
      - 'QN. text' or 'N. text'
    Answer line: 'Answer: a' or 'Ans: (b)' or 'Ans.(C)'
    Options:     '(a) text'  '(b) text'  etc.
    """
    # Normalize: collapse runs of spaces, strip trailing spaces per line
    lines = [re.sub(r" {2,}", " ", ln).strip() for ln in text.splitlines()]
    text = "\n".join(lines)

    # ---- Locate every question block -------------------------------------------
    # Matches: "Question: 5." OR "Q5." OR standalone "5."
    q_start = re.compile(
        r'(?:Question:\s*(\d+)\.|(?:Q\.?\s*)?(\d+)[.)]\s)',
        re.IGNORECASE,
    )

    # Find all start positions + question numbers
    matches = list(q_start.finditer(text))
    if not matches:
        return []

    blocks = []
    for i, m in enumerate(matches):
        q_num = m.group(1) or m.group(2)
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        blocks.append({"num": q_num, "body": body})

    # ---- For each block, parse stem, options, answer ---------------------------
    opt_pattern = re.compile(
        r'\(\s*([A-Da-d])\s*\)\s*(.+?)(?=\(\s*[A-Da-d]\s*\)|Answer:|Ans[.:]|Positive|$)',
        re.DOTALL | re.IGNORECASE,
    )
    ans_pattern = re.compile(
        r'(?:Answer|Ans)[.:\s]*\(?\s*([A-Da-d])\s*\)?',
        re.IGNORECASE,
    )

    parsed = []
    for blk in blocks:
        body = blk["body"]

        # Extract options
        options: dict[str, str] = {}
        for om in opt_pattern.finditer(body):
            key = om.group(1).upper()
            val = om.group(2).strip().split("\n")[0].strip()
            # Clean trailing noise
            val = re.split(r'\s{2,}', val)[0].strip()
            options[key] = val

        # Extract answer letter
        ans_m = ans_pattern.search(body)
        answer_letter = ans_m.group(1).upper() if ans_m else None
        answer_text = options.get(answer_letter, answer_letter) if answer_letter else "N/A"

        # Extract question stem (text before first option)
        first_opt = opt_pattern.search(body)
        if first_opt:
            stem = body[: first_opt.start()].strip()
        else:
            stem = ans_pattern.sub("", body).strip()
            stem = stem.split("\n")[0].strip()

        # Clean up stem: remove trailing punctuation noise, "Positive Marks" etc.
        stem = re.split(r'Positive Marks|Negative Marks', stem)[0].strip()

        parsed.append({
            "num": blk["num"],
            "question": stem,
            "options": options,
            "answer_letter": answer_letter,
            "answer_text": answer_text,
        })

    return parsed


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def format_answer_line(item: dict) -> str:
    letter = item["answer_letter"] or ""
    text = item["answer_text"]
    if not letter or letter == text:
        return f"Answer - {text}"
    return f"Answer - {text}"  # show only the resolved text, not the letter


# ---------------------------------------------------------------------------
# Output generators
# ---------------------------------------------------------------------------

def _safe(text: str) -> str:
    """Replace characters outside latin-1 with '?' so fpdf2 doesn't crash."""
    return text.encode("latin-1", errors="replace").decode("latin-1")


def generate_pdf(qa_list: list[dict], title: str) -> str:
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, _safe(title), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(5)

    for item in qa_list:
        q_line = _safe(re.sub(r'\s+', ' ', f"Q{item['num']}. {item['question']}").strip())
        a_line = _safe(re.sub(r'\s+', ' ', format_answer_line(item)).strip())

        pdf.set_font("Helvetica", "B", 11)
        pdf.multi_cell(0, 7, q_line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 7, a_line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(3)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf", prefix="answers_")
    pdf.output(tmp.name)
    return tmp.name


def generate_docx(qa_list: list[dict], title: str) -> str:
    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for item in qa_list:
        q_line = re.sub(r'\s+', ' ', f"Q{item['num']}. {item['question']}").strip()
        a_line = re.sub(r'\s+', ' ', format_answer_line(item)).strip()

        p_q = doc.add_paragraph()
        run = p_q.add_run(q_line)
        run.bold = True
        run.font.size = Pt(11)

        p_a = doc.add_paragraph()
        ar = p_a.add_run(a_line)
        ar.font.size = Pt(11)

        doc.add_paragraph()  # blank spacer

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", prefix="answers_")
    doc.save(tmp.name)
    return tmp.name


# ---------------------------------------------------------------------------
# Gradio handler
# ---------------------------------------------------------------------------

def process(pdf_file, output_format: str, doc_title: str):
    if pdf_file is None:
        return "Please upload a PDF file.", None

    try:
        raw_text = extract_text_column_aware(pdf_file)
    except Exception as e:
        return f"Error reading PDF: {e}", None

    if not raw_text.strip():
        return "Could not extract text. This may be a scanned image PDF (not supported).", None

    qa_list = parse_qa(raw_text)

    if not qa_list:
        preview = "No questions detected.\n\nRaw extracted text (first 2000 chars):\n\n" + raw_text[:2000]
        return preview, None

    # Build preview
    lines = []
    for item in qa_list:
        lines.append(f"Q{item['num']}. {item['question']}")
        lines.append(format_answer_line(item))
        lines.append("")
    preview = "\n".join(lines)

    title = doc_title.strip() or "Answer Sheet"
    if output_format == "PDF":
        out_path = generate_pdf(qa_list, title)
    else:
        out_path = generate_docx(qa_list, title)

    return preview, out_path


# ---------------------------------------------------------------------------
# Gradio UI
# ---------------------------------------------------------------------------

with gr.Blocks(title="MCQ Answer Sheet Generator") as demo:
    gr.Markdown("# MCQ PDF → Answer Sheet Generator")
    gr.Markdown(
        "Upload a PDF with MCQ questions. The app extracts each question with its correct answer "
        "and generates a clean, readable answer sheet."
    )

    with gr.Row():
        with gr.Column(scale=1):
            pdf_input = gr.File(
                label="Upload MCQ PDF",
                file_types=[".pdf"],
            )
            title_input = gr.Textbox(
                label="Document Title",
                value="Answer Sheet",
                placeholder="e.g. GK Weekly Test 11 – Answers",
            )
            format_input = gr.Radio(
                choices=["PDF", "DOCX"],
                label="Output Format",
                value="PDF",
            )
            submit_btn = gr.Button("Generate Answer Sheet", variant="primary", size="lg")

        with gr.Column(scale=2):
            preview_output = gr.Textbox(
                label="Preview",
                lines=35,
                interactive=False,
                placeholder="Q&A preview will appear here after processing...",
            )
            file_output = gr.File(label="Download Answer Sheet")

    submit_btn.click(
        fn=process,
        inputs=[pdf_input, format_input, title_input],
        outputs=[preview_output, file_output],
    )

if __name__ == "__main__":
    demo.launch(share=True)
